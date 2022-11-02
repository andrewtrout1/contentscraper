from bs4 import BeautifulSoup
from urllib.request import Request, urlopen
from docx import Document



# Place URLs to scrape here (wrapped in quotes and separated by a comma)
urls = []


headers = {
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'GET',
    'Access-Control-Allow-Headers': 'Content-Type',
    'Access-Control-Max-Age': '3600',
    'User-Agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:52.0) Gecko/20100101 Firefox/52.0'
}

for url in urls:
    try:
        req = Request(url, headers=headers)
        webpage = urlopen(req)
        soup = BeautifulSoup(webpage, 'html.parser')
        # print(url)
        title = str(soup.find("h1", {"class": "top-page-title"}).get_text())
        # print(title)
        content = str(soup.find("div", {"class": "entry-text"}).get_text())
        # print(content)

        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph(content)

        document.save(title + '.docx')

        

    except:
        print(url + " - Error!!")